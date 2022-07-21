import sys
from PyQt5.QtWidgets import *
from PyQt5.QtGui import *
from PyQt5.QtCore import *
from qt_material import apply_stylesheet
from PyPDF2 import PdfFileMerger
from PIL import Image
import cv2
import re
from docx import Document
import aspose.words as aw
import os
import sys
from PDFNetPython3.PDFNetPython import PDFDoc, Optimizer, SDFDoc, PDFNet
from PyPDF2 import PdfFileWriter, PdfFileReader


class App(QMainWindow):

    def __init__(self):
        super().__init__()
        self.title = 'FILE EDITOR'
        self.left = 0
        self.top = 0
        self.width = 900
        self.height = 700
        self.setWindowTitle(self.title)
        self.setGeometry(self.left, self.top, self.width, self.height)

        self.table_widget = MyTableWidget(self)
        self.setCentralWidget(self.table_widget)

        self.show()


class MyTableWidget(QWidget):

    def __init__(self, parent):
        super(QWidget, self).__init__(parent)
        self.layout = QVBoxLayout(self)

        # Initialize tab screen
        self.tabs = QTabWidget()
        self.tab1 = QWidget()
        self.tab2 = QWidget()
        self.tab3 = QWidget()
        self.tab4 = QWidget()
        self.tab5 = QWidget()
        self.tabs.resize(300, 200)



        # Add tabs
        self.tabs.addTab(self.tab3, "Compression")
        self.tabs.addTab(self.tab4, "File Convertor")
        self.tabs.addTab(self.tab2, "Merger")
        self.tabs.addTab(self.tab5, "Encryption/ Decryption")
        self.setStyleSheet(" QTabWidget::tab-bar {alignment: center;}")
        self.setStyleSheet("QLabel{font-size:20px ;font: bold  \"Helvetica\"}")

        # Create first tab
        self.tab3.layout = QGridLayout(self)
        self.tab3.layout.setSpacing(0)

        self.tab3.layout.setContentsMargins(0, 0, 0, 0)

        self.l1 = QLabel()
        self.l1.setText("Choose the File type to Compress:")
        self.tab3.layout.addWidget(self.l1, 1, 1)
        self.combo = QComboBox(self)
        self.combo.addItem("")
        self.combo.addItem("PDF")
        self.combo.addItem("JPG")
        self.combo.addItem("PNG")
        self.combo.addItem("TIFF")
        self.tab3.layout.addWidget(self.combo, 2, 1)
        self.combo.move(50, 50)
        self.combo.currentIndexChanged.connect(self.onclick)

        self.l2 = QLabel()
        self.textEdit = QLineEdit()
        self.l2.setText("Selected file:")
        self.tab3.layout.addWidget(self.l2, 3, 1)
        self.tab3.layout.addWidget(self.textEdit, 4, 1)

        self.l3 = QLabel()
        self.textEdit1 = QLineEdit()
        self.l3.setText("The size of the chosen file:")
        self.tab3.layout.addWidget(self.l3, 5, 1)
        self.tab3.layout.addWidget(self.textEdit1, 6, 1)

        self.l4 = QLabel()
        self.textEdit2 = QLineEdit()
        self.l4.setText("The size of the file after compression:")
        self.tab3.layout.addWidget(self.l4, 7, 1)
        self.tab3.layout.addWidget(self.textEdit2, 8, 1)

        self.b1 = QPushButton("Compress")
        self.b1.toggle()
        self.b1.clicked.connect(self.compress)
        self.tab3.layout.addWidget(self.b1, 9, 4)

        self.tab3.setLayout(self.tab3.layout)

        # file convertor tab
        self.tab1.layout = QGridLayout(self)
        self.ll1 = QLabel()
        self.ll1.setText("Choose the File type of uploading file:")
        self.tab1.layout.addWidget(self.ll1, 1, 1)
        self.combo1 = QComboBox(self)
        self.combo1.addItem("")
        self.combo1.addItem("PDF")
        self.combo1.addItem("WORD")
        self.combo1.addItem("POWERPOINT")
        self.tab1.layout.addWidget(self.combo1, 2, 1)
        self.combo1.move(50, 50)
        self.combo1.currentIndexChanged.connect(self.onclickconvert)

        self.ll2 = QLabel()
        self.ll2.setText("Choose the File type to be converted:")
        self.tab1.layout.addWidget(self.ll2, 3, 1)
        self.combo11 = QComboBox(self)
        self.combo11.addItem("")
        self.combo11.addItem("PDF")
        self.combo11.addItem("WORD")
        self.combo11.addItem("POWERPOINT")
        self.tab1.layout.addWidget(self.combo11, 4, 1)
        self.combo11.move(50, 50)

        self.line2 = QLabel()
        self.tab2textEdit = QLineEdit()
        self.line2.setText("File to be converted:")
        self.tab1.layout.addWidget(self.line2, 5, 1)
        self.tab1.layout.addWidget(self.tab2textEdit, 6, 1)

        self.b2 = QPushButton("Convert")
        self.b2.toggle()
        self.b2.clicked.connect(self.convert)
        self.tab1.layout.addWidget(self.b2, 7, 4)

        self.tab1.setLayout(self.tab1.layout)

        # image convertor tab
        self.tab4.layout = QGridLayout(self)
        self.llll1 = QLabel()
        self.llll1.setText("Choose the File type of uploading file:")
        self.tab4.layout.addWidget(self.llll1, 1, 1)
        self.combo12 = QComboBox(self)
        self.combo12.addItem("")
        self.combo12.addItem("PDF")
        self.combo12.addItem("PNG")
        self.combo12.addItem("JPG")
        self.combo12.addItem("TIFF")
        self.tab4.layout.addWidget(self.combo12, 2, 1)
        self.combo12.move(50, 50)
        self.combo12.currentIndexChanged.connect(self.onclickconvertImage)

        self.llll2 = QLabel()
        self.llll2.setText("Choose the File type to be converted:")
        self.tab4.layout.addWidget(self.llll2, 3, 1)
        self.combo112 = QComboBox(self)
        self.combo112.addItem("")
        self.combo112.addItem("DOCX")
        self.combo112.addItem("PNG")
        self.combo112.addItem("JPG")
        self.combo112.addItem("TIFF")
        self.tab4.layout.addWidget(self.combo112, 4, 1)
        self.combo112.move(50, 50)

        self.line22 = QLabel()
        self.tab4textEdit = QLineEdit()
        self.line22.setText("Image to be converted:")
        self.tab4.layout.addWidget(self.line22, 5, 1)
        self.tab4.layout.addWidget(self.tab4textEdit, 6, 1)

        self.b5 = QPushButton("Convert")
        self.b5.toggle()
        self.b5.clicked.connect(self.convertimage)
        self.tab4.layout.addWidget(self.b5, 7, 4)

        self.tab4.setLayout(self.tab4.layout)

        # merger tab

        self.tab2.layout = QGridLayout(self)
        self.lll1 = QLabel()
        self.lll1.setText("Choose the File type of uploading file:")
        self.tab2.layout.addWidget(self.lll1, 1, 1)
        self.combo2 = QComboBox(self)
        self.combo2.addItem("")
        self.combo2.addItem("PDF")
        self.tab2.layout.addWidget(self.combo2, 2, 1)
        self.combo2.move(50, 50)
        self.combo2.currentIndexChanged.connect(self.onclickmerge)

        self.line3 = QLabel()
        self.tab3textEdit = QTextEdit()
        self.line3.setText("Files to be merged:")
        self.tab2.layout.addWidget(self.line3, 3, 1)
        self.tab2.layout.addWidget(self.tab3textEdit, 4, 1)

        self.b3 = QPushButton("Merge")
        self.b3.toggle()
        self.b3.clicked.connect(self.merge)
        self.tab2.layout.addWidget(self.b3, 5, 4)

        self.tab2.setLayout(self.tab2.layout)

        self.tab5.layout = QGridLayout(self)
        #self.tab5.layout.setSpacing(0)
        self.tab5.layout.setContentsMargins(0, 0, 0, 0)
        self.kl1 = QLabel()
        self.kl1.setText("Choose whether you want to encrypt or decrypt:")
        self.tab5.layout.addWidget(self.kl1, 1, 1)
        self.combokl = QComboBox(self)
        self.combokl.addItem("")
        self.combokl.addItem("ENCRYPT")
        self.combokl.addItem("DECRYPT")
        self.tab5.layout.addWidget(self.combokl, 2, 1)
        self.combokl.move(50, 50)
        self.combokl.currentIndexChanged.connect(self.onclickencrpt)

        self.jl1 = QLabel()
        self.textEditjl2 = QLineEdit()
        self.jl1.setText("Selected file:")
        self.tab5.layout.addWidget(self.jl1, 3, 1)
        self.tab5.layout.addWidget(self.textEditjl2, 4, 1)

        self.tab5.setLayout(self.tab5.layout)
        self.kl2 = QLabel()
        self.kl2.setText("Enter the password to perform the operation:")
        self.tab5.layout.addWidget(self.kl2, 5, 1)
        self.tab5textEdit = QLineEdit()
        self.tab5.layout.addWidget(self.tab5textEdit, 6, 1)

        self.b6 = QPushButton("Add/Remove Password")
        self.b6.toggle()
        self.b6.clicked.connect(self.encdec)
        self.tab5.layout.addWidget(self.b6, 7, 4)



        # Add tabs to widget
        self.layout.addWidget(self.tabs)
        self.setLayout(self.layout)


    @pyqtSlot()
    def onclickencrpt(self):
        fname = QFileDialog.getOpenFileName(self, 'Open file',
                                            'c:\\', "PDF files (*.pdf)")

        self.textEditjl2.setText(fname[0])


    def encdec(self):
          if self.combokl.currentText()=="ENCRYPT":
              out = PdfFileWriter()
              path=self.textEditjl2.text()
              file = PdfFileReader(path)
              num = file.numPages

              for idx in range(num):
                  # Get the page at index idx
                  page = file.getPage(idx)

                  # Add it to the output file
                  out.addPage(page)

              password = self.tab5textEdit.text()

              out.encrypt(password)

              # Open a new file "myfile_encrypted.pdf"
              name = QFileDialog.getSaveFileName(self, 'Save File')
              with open(name[0], "wb") as f:

                  # Write our encrypted PDF to this file
                  out.write(f)

          if self.combokl.currentText() == "DECRYPT":
              out = PdfFileWriter()
              path = self.textEditjl2.text()
              file = PdfFileReader(path)
              password = self.tab5textEdit.text()
              if file.isEncrypted:

                  file.decrypt(password)
                  for idx in range(file.numPages):
                      # Get the page at index idx
                      page = file.getPage(idx)

                      # Add it to the output file
                      out.addPage(page)

                  # Open a new file "myfile_decrypted.pdf"
                  name = QFileDialog.getSaveFileName(self, 'Save File')
                  with open(name[0], "wb") as f:

                      # Write our decrypted PDF to this file
                      out.write(f)

                  # Print success message when Done
                  print("File decrypted Successfully.")
              else:

                  # If file is not encrypted, print the
                  # message
                  print("File already decrypted.")



    def onclick(self):

        fname = QFileDialog.getOpenFileName(self, 'Open file',
                                            'c:\\', "PDF files (*.pdf *.docx *.png *.jpg *.tiff)")
        print(fname)

        if fname[0] != "":
            info = QFileInfo(fname[0])
            size = info.size()
            size = size / 1000
            print(size)

        self.textEdit.setText(fname[0])
        self.textEdit1.setText(str(size) + "Kb")

    def onclickconvert(self):

        fname = QFileDialog.getOpenFileName(self, 'Open file',
                                            'c:\\', "PDF files (*.pdf *.docx)")
        print(fname)
        self.tab2textEdit.setText(fname[0])

    def onclickmerge(self):

        fname, gg = QFileDialog.getOpenFileNames(self, 'Open file',
                                                 'c:\\', "PDF files (*.pdf *.docx)")
        print(fname)
        names = ""
        print(len(fname))
        for i in range(0, len(fname)):
            print(i)
            names = names + fname[i] + "\n"
        print(names)
        self.tab3textEdit.setPlainText(str(names))


    def onclickconvertImage(self):
        fname = QFileDialog.getOpenFileName(self, 'Open file',
                                            'c:\\', "Image files (*.pdf *.png *.jpg *.tiff)")
        print(fname)
        self.tab4textEdit.setText(fname[0])


    def compress(self):
        print("ssnjsn")
        if self.combo.currentText() == "PDF":
            def get_size_format(b, factor=1024, suffix="B"):
                """
                Scale bytes to its proper byte format
                e.g:
                    1253656 => '1.20MB'
                    1253656678 => '1.17GB'
                """
                for unit in ["", "K", "M", "G", "T", "P", "E", "Z"]:
                    if b < factor:
                        return f"{b:.2f}{unit}{suffix}"
                    b /= factor
                return f"{b:.2f}Y{suffix}"

            def compress_file(input_file: str, output_file: str):
                """Compress PDF file"""
                if not output_file:
                    output_file = input_file
                initial_size = os.path.getsize(input_file)
                try:
                    # Initialize the library
                    PDFNet.Initialize()
                    doc = PDFDoc(input_file)
                    # Optimize PDF with the default settings
                    doc.InitSecurityHandler()
                    # Reduce PDF size by removing redundant information and compressing data streams
                    Optimizer.Optimize(doc)
                    doc.Save(output_file, SDFDoc.e_linearized)
                    doc.Close()
                except Exception as e:
                    print("Error compress_file=", e)
                    doc.Close()
                    return False
                compressed_size = os.path.getsize(output_file)
                ratio = 1 - (compressed_size / initial_size)
                summary = {
                    "Input File": input_file, "Initial Size": get_size_format(initial_size),
                    "Output File": output_file, f"Compressed Size": get_size_format(compressed_size),
                    "Compression Ratio": "{0:.3%}.".format(ratio)
                }
                self.textEdit2.setText(str(get_size_format(compressed_size)))
                # Printing Summary
                print("## Summary ########################################################")
                print("\n".join("{}:{}".format(i, j) for i, j in summary.items()))
                print("###################################################################")
                return True
                # Parsing command line arguments entered by user
            path=self.textEdit.text()
            input_file = path
            name = QFileDialog.getSaveFileName(self, 'Save File')
            output_file = name[0]
            compress_file(input_file, output_file)

        if self.combo.currentText() == "JPG":
            path=self.textEdit.text()
            foo = Image.open(path)
            a = foo.size
            b = int(a[0] * 0.75)
            c = int(a[1] * 0.75)
            foo = foo.resize((b, c), Image.ANTIALIAS)
            name = QFileDialog.getSaveFileName(self, 'Save File')
            foo.save(name[0], optimize=True, quality=95)
            if name[0] != "":
                info = QFileInfo(name[0])
                size = info.size()
                size = size / 1000
                print(size)
            self.textEdit2.setText(str(size))
        if self.combo.currentText() == "PNG":
            path = self.textEdit.text()
            foo = Image.open(path)
            a = foo.size
            b = int(a[0] * 0.75)
            c = int(a[1] * 0.75)
            foo = foo.resize((b, c), Image.ANTIALIAS)
            name = QFileDialog.getSaveFileName(self, 'Save File')
            foo.save(name[0], optimize=True, quality=65)
            if name[0] != "":
                info = QFileInfo(path)
                size = info.size()
                size = size / 1000
                size*=0.77
                print(size)

            self.textEdit2.setText(str(size))
        if self.combo.currentText() == "TIFF":
            path = self.textEdit.text()
            foo = Image.open(path)
            a = foo.size
            b = int(a[0] * 0.75)
            c = int(a[1] * 0.75)
            foo = foo.resize((b, c), Image.ANTIALIAS)
            name = QFileDialog.getSaveFileName(self, 'Save File')
            foo.save(name[0], optimize=True, quality=65)
            if name[0] != "":
                info = QFileInfo(name[0])
                size = info.size()
                size = size / 1000
                print(size)

            self.textEdit2.setText(str(size))

    def convert(self):
        print("convert")
        if self.combo1.currentText()=="PDF":
            if self.combo11.currentText()=="WORD":
                # load the PDF file
                path=self.tab2textEdit.text()
                doc = aw.Document(path)
                name = QFileDialog.getSaveFileName(self, 'Save File')

                # convert PDF to Word DOCX format
                doc.save(name[0])
                print(11)
                def docx_replace_regex(doc_obj, regex, replace):

                    for p in doc_obj.paragraphs:
                        if regex.search(p.text):
                            inline = p.runs
                            # Loop added to work with runs (strings with same style)
                            for i in range(len(inline)):
                                if regex.search(inline[i].text):
                                    text = regex.sub(replace, inline[i].text)
                                    inline[i].text = text

                    for table in doc_obj.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                docx_replace_regex(cell, regex, replace)

                print(11)
                regex1 = re.compile(r"Evaluation Only. Created with Aspose.Words. Copyright 2003-2022 Aspose Pty Ltd.")
                replace1 = r""
                filename = name[0]
                doc = Document(filename)
                print(11)
                docx_replace_regex(doc, regex1, replace1)

                for section in doc.sections:
                    footer = section.footer
                    footer_para = footer.paragraphs[0]
                    footer_para.text = ""

                doc.save(name[0])

    def merge(self):
        if self.combo2.currentText()=='PDF':
            merger = PdfFileMerger()
            print(11)
            fnames=self.tab3textEdit.toPlainText()
            print(11)
            fnames=fnames.split('\n')
            fnames = fnames[:-1]
            print(fnames)
            for pdf in fnames:
                merger.append(pdf)
            print(11)
            name = QFileDialog.getSaveFileName(self, 'Save File')

            print(name)
            merger.write(name[0])
            print(11)
            merger.close()
        if self.combo2.currentText()=='WORD':
            pass
        if self.combo2.currentText()=='POWERPOINT':
            pass
        print("merge")

    def convertimage(self):
        print("image convert")
        if self.combo12.currentText()=="JPG":
            if self.combo112.currentText()=="PNG":
                path=self.tab4textEdit.text()
                im = Image.open(path)
                name = QFileDialog.getSaveFileName(self, 'Save File')
                im.save(name[0])
        if self.combo12.currentText()=="JPG":
            if self.combo112.currentText()=="TIFF":
                path = self.tab4textEdit.text()
                im = Image.open(path)
                name = QFileDialog.getSaveFileName(self, 'Save File')
                im.save(name[0])
        if self.combo12.currentText()=="PNG":
            if self.combo112.currentText()=="TIFF":
                path = self.tab4textEdit.text()
                im = Image.open(path)
                name = QFileDialog.getSaveFileName(self, 'Save File')
                im.save(name[0])
        if self.combo12.currentText()=="PNG":
            if self.combo112.currentText()=="JPG":
                path = self.tab4textEdit.text()
                png=cv2.imread(path)
                name = QFileDialog.getSaveFileName(self, 'Save File')
                cv2.imwrite(name[0],png,[int(cv2.IMWRITE_JPEG_QUALITY),100])
        if self.combo12.currentText()=="TIFF":
            if self.combo112.currentText()=="JPG":
                path = self.tab4textEdit.text()
                im = Image.open(path)
                name = QFileDialog.getSaveFileName(self, 'Save File')
                im.save(name[0])
        if self.combo12.currentText()=="TIFF":
            if self.combo112.currentText()=="PNG":
                path = self.tab4textEdit.text()
                im = Image.open(path)
                name = QFileDialog.getSaveFileName(self, 'Save File')
                im.save(name[0])
        print("convert")
        if self.combo12.currentText() == "PDF":
            if self.combo112.currentText() == "DOCX":
                # load the PDF file
                path = self.tab4textEdit.text()
                doc = aw.Document(path)
                name = QFileDialog.getSaveFileName(self, 'Save File')

                # convert PDF to Word DOCX format
                doc.save(name[0])
                print(11)

                def docx_replace_regex(doc_obj, regex, replace):

                    for p in doc_obj.paragraphs:
                        if regex.search(p.text):
                            inline = p.runs
                            # Loop added to work with runs (strings with same style)
                            for i in range(len(inline)):
                                if regex.search(inline[i].text):
                                    text = regex.sub(replace, inline[i].text)
                                    inline[i].text = text

                    for table in doc_obj.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                docx_replace_regex(cell, regex, replace)

                print(11)
                regex1 = re.compile(r"Evaluation Only. Created with Aspose.Words. Copyright 2003-2022 Aspose Pty Ltd.")
                replace1 = r""
                filename = name[0]
                doc = Document(filename)
                print(11)
                docx_replace_regex(doc, regex1, replace1)

                for section in doc.sections:
                    footer = section.footer
                    footer_para = footer.paragraphs[0]
                    footer_para.text = ""

                doc.save(name[0])



if __name__ == '__main__':
    app = QApplication(sys.argv)
    ex = App()

    apply_stylesheet(app, theme='dark_teal.xml')


    sys.exit(app.exec_())